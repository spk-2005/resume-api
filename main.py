"""
main.py  ─  Resume ATS Intelligence API  v3.0.0
"""

import io
import os
import logging
import tempfile
import traceback
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Optional

from fastapi import (
    FastAPI, APIRouter, Depends, HTTPException,
    UploadFile, File, Form, Request, status
)
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordRequestForm
from sqlalchemy import extract
from sqlalchemy.orm import Session

import requests
import re
from database import SessionLocal, engine, get_db
import models, schemas
from services.resume_analyzer import analyze_resume
from services.auth import create_access_token, get_current_user, get_password_hash, verify_password, ACCESS_TOKEN_EXPIRE_MINUTES
from fastapi.concurrency import run_in_threadpool


# ══════════════════════════════════════════════════════════════════
#  LOGGING  — errors will now appear in your server logs
# ══════════════════════════════════════════════════════════════════

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(message)s",
)
logger = logging.getLogger(__name__)


# ══════════════════════════════════════════════════════════════════
#  OPTIONAL IMPORTS
# ══════════════════════════════════════════════════════════════════

try:
    from pypdf import PdfReader
    PYPDF_OK = True
except ImportError:
    PYPDF_OK = False
    logger.warning("pypdf not installed — PDF support disabled.")

try:
    from pdf2image import convert_from_bytes
    import pytesseract
    OCR_OK = True
except ImportError:
    OCR_OK = False
    logger.warning("pdf2image/pytesseract not installed — scanned PDF OCR disabled.")

try:
    import docx
    DOCX_OK = True
except ImportError:
    DOCX_OK = False
    logger.warning("python-docx not installed — DOCX support disabled.")

try:
    import textract
    TEXTRACT_OK = True
except ImportError:
    TEXTRACT_OK = False
    logger.warning("textract not installed — DOC support disabled.")

try:
    from PIL import Image
    import pytesseract
    PIL_OK = True
except ImportError:
    PIL_OK = False
    logger.warning("Pillow/pytesseract not installed — image OCR disabled.")


# ══════════════════════════════════════════════════════════════════
#  CONFIG
# ══════════════════════════════════════════════════════════════════

PDF_EXTS   = {".pdf"}
DOCX_EXTS  = {".docx"}
DOC_EXTS   = {".doc"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tiff"}
TEXT_EXTS  = {".txt", ".md", ".rtf", ".csv"}
HTML_EXTS  = {".html", ".htm"}
ALL_SUPPORTED = PDF_EXTS | DOCX_EXTS | DOC_EXTS | IMAGE_EXTS | TEXT_EXTS | HTML_EXTS


# ══════════════════════════════════════════════════════════════════
#  APP
# ══════════════════════════════════════════════════════════════════

app = FastAPI(
    title="Resume ATS Intelligence API",
    version="3.0.0",
)

# Enable CORS for frontend development and production
_cors_origins = os.getenv(
    "CORS_ORIGINS",
    "http://localhost:3000,http://localhost:5173,http://127.0.0.1:5173",
)
app.add_middleware(
    CORSMiddleware,
    allow_origins=[origin.strip() for origin in _cors_origins.split(",") if origin.strip()],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# Create all DB tables on startup
try:
    models.Base.metadata.create_all(bind=engine)
    logger.info("Database tables created/verified OK.")
except Exception as e:
    logger.error(f"FATAL: Could not create DB tables: {e}")


# ══════════════════════════════════════════════════════════════════
#  HEALTH CHECK  — required for RapidAPI health check to pass
#  Set Health Check URL to: https://resume-api-c908.onrender.com/health
# ══════════════════════════════════════════════════════════════════

@app.get("/health", tags=["Health"], include_in_schema=False)
@app.get("/",       tags=["Health"], include_in_schema=False)
async def health_check():
    """
    Health check endpoint.
    RapidAPI pings this daily — must return HTTP 200.
    Also handles root URL so there is no 404.
    """
    db_status = "ok"
    try:
        db = SessionLocal()
        db.execute(__import__("sqlalchemy").text("SELECT 1"))
        db.close()
    except Exception as e:
        db_status = f"error: {e}"

    return {
        "status":  "ok",
        "service": "Resume ATS Intelligence API",
        "version": "3.0.0",
        "database": db_status,
    }


# ══════════════════════════════════════════════════════════════════
#  GLOBAL ERROR HANDLER  — turns any unhandled crash into clean JSON
# ══════════════════════════════════════════════════════════════════

@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    tb = traceback.format_exc()
    logger.error(f"Unhandled error on {request.url.path}:\n{tb}")
    return JSONResponse(
        status_code=500,
        content={
            "error":   "Internal Server Error",
            "detail":  str(exc),
            "type":    type(exc).__name__,
            # Remove 'trace' in production for security
            "trace":   tb.splitlines()[-3:],
        },
    )


# ══════════════════════════════════════════════════════════════════
#  DB
# ══════════════════════════════════════════════════════════════════

def _ats_score_value(report: dict) -> float:
    """Extract numeric ATS score from analyzer report."""
    scores = report.get("scores") or {}
    ats = scores.get("ats_score") or {}
    return float(ats.get("value") or 0)


# ══════════════════════════════════════════════════════════════════
#  USAGE TRACKING
# ══════════════════════════════════════════════════════════════════

def _count_usage_this_month(user_id: int, db: Session) -> int:
    now = datetime.utcnow()
    # Range-based query is more index-friendly than extract()
    start_of_month = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    try:
        return db.query(models.UsageLog).filter(
            models.UsageLog.user_id == user_id,
            models.UsageLog.request_time >= start_of_month
        ).count()
    except Exception as e:
        logger.error(f"Error counting usage for user {user_id}: {e}", exc_info=True)
        return 0


def log_usage(user_id: int, endpoint: str, db: Session):
    try:
        db.add(models.UsageLog(
            user_id=user_id,
            request_time=datetime.utcnow(),
            endpoint=endpoint,
            status="success",
        ))
        db.commit()
        logger.info(f"Usage logged: user_id={user_id} endpoint={endpoint}")
    except Exception as e:
        logger.error(f"Failed to log usage for user {user_id}: {e}")
        db.rollback()


# ══════════════════════════════════════════════════════════════════
#  AUTH ENDPOINTS
# ══════════════════════════════════════════════════════════════════

auth_router = APIRouter(tags=["Authentication"])

@auth_router.post("/register", response_model=schemas.UserResponse)
def register_user(user: schemas.UserCreate, db: Session = Depends(get_db)):
    db_user = db.query(models.User).filter(models.User.email == user.email).first()
    if db_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    hashed_password = get_password_hash(user.password)
    new_user = models.User(email=user.email, hashed_password=hashed_password) # Assumes User model has hashed_password
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    return schemas.UserResponse.model_validate(new_user)


@auth_router.post("/login", response_model=schemas.Token)
def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    user = db.query(models.User).filter(models.User.email == form_data.username).first()
    if not user or not verify_password(form_data.password, user.hashed_password):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect email or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user.email}, expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer"}


@app.get("/users/me", response_model=schemas.UserResponse, tags=["Users"])
def read_users_me(current_user: models.User = Depends(get_current_user)):
    return schemas.UserResponse.model_validate(current_user)

app.include_router(auth_router, prefix="/auth")

# ══════════════════════════════════════════════════════════════════
#  TEXT EXTRACTION
# ══════════════════════════════════════════════════════════════════

def _extract_pdf(file_bytes: bytes, filename: str) -> str:
    if not PYPDF_OK:
        raise HTTPException(status_code=501, detail="pypdf not installed. Run: pip install pypdf")
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text   = "\n".join(p.extract_text() or "" for p in reader.pages).strip()
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Could not read PDF '{filename}': {e}")

    if text:
        return text

    if not OCR_OK:
        raise HTTPException(status_code=422, detail=f"'{filename}' is a scanned PDF and OCR is not available.")
    try:
        images = convert_from_bytes(file_bytes, dpi=200)
        text   = "\n".join(pytesseract.image_to_string(img, lang="eng") for img in images).strip()
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"OCR failed on '{filename}': {e}")

    if not text:
        raise HTTPException(status_code=422, detail=f"No text extracted from '{filename}' even after OCR.")
    return text


def _extract_docx(file_bytes: bytes, filename: str) -> str:
    if not DOCX_OK:
        raise HTTPException(status_code=501, detail="python-docx not installed.")
    try:
        document = docx.Document(io.BytesIO(file_bytes))
        parts    = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        for section in document.sections:
            for p in (section.header.paragraphs if section.header else []):
                parts.append(p.text)
            for p in (section.footer.paragraphs if section.footer else []):
                parts.append(p.text)
        text = "\n".join(p for p in parts if p.strip()).strip()
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Could not read DOCX '{filename}': {e}")
    if not text:
        raise HTTPException(status_code=422, detail=f"No text extracted from '{filename}'.")
    return text


def _extract_doc(file_bytes: bytes, filename: str) -> str:
    if not TEXTRACT_OK:
        raise HTTPException(status_code=501, detail="textract not installed.")
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        text = textract.process(tmp_path).decode("utf-8", errors="replace").strip()
        os.unlink(tmp_path)
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Could not read DOC '{filename}': {e}")
    if not text:
        raise HTTPException(status_code=422, detail=f"No text extracted from '{filename}'.")
    return text


def _extract_image(file_bytes: bytes, filename: str) -> str:
    if not PIL_OK:
        raise HTTPException(status_code=501, detail="Pillow/pytesseract not installed.")
    try:
        text = pytesseract.image_to_string(Image.open(io.BytesIO(file_bytes)), lang="eng").strip()
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"OCR failed on '{filename}': {e}")
    if not text:
        raise HTTPException(status_code=422, detail=f"No text found in image '{filename}'.")
    return text


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext in TEXT_EXTS:   return file_bytes.decode("utf-8", errors="replace").strip()
    if ext in PDF_EXTS:    return _extract_pdf(file_bytes, filename)
    if ext in DOCX_EXTS:   return _extract_docx(file_bytes, filename)
    if ext in DOC_EXTS:    return _extract_doc(file_bytes, filename)
    if ext in IMAGE_EXTS:  return _extract_image(file_bytes, filename)
    
    # Generic fallback using textract if available
    if TEXTRACT_OK:
        try:
            with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            text = textract.process(tmp_path).decode("utf-8", errors="replace").strip()
            os.unlink(tmp_path)
            if text: return text
        except Exception as e:
            logger.warning(f"Textract fallback failed for {filename}: {e}")
            
    if ext in HTML_EXTS:
        text = file_bytes.decode("utf-8", errors="replace")
        text = re.sub(r'<[^>]+>', ' ', text)
        return re.sub(r'\s+', ' ', text).strip()

    raise HTTPException(
        status_code=415,
        detail=f"Unsupported file type '{ext}'. Accepted: {', '.join(sorted(ALL_SUPPORTED))}",
    )


def _extract_from_url(url: str) -> str:
    """Scrapes text from a job posting URL."""
    try:
        logger.info(f"Scraping JD from URL: {url}")
        # Using a browser-like User-Agent to avoid some basic anti-bot measures
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        }
        response = requests.get(url, timeout=15, headers=headers)
        response.raise_for_status()
        
        # Simple HTML extraction
        html = response.text
        # Remove scripts and styles
        html = re.sub(r'<(script|style|header|footer|nav)[^>]*>.*?</\1>', ' ', html, flags=re.DOTALL | re.IGNORECASE)
        # Remove all other tags
        text = re.sub(r'<[^>]+>', ' ', html)
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        if len(text) < 50:
            raise ValueError("Extracted text too short (likely blocked or empty page).")
        return text
    except Exception as e:
        logger.error(f"Failed to scrape URL {url}: {e}")
        raise HTTPException(
            status_code=422, 
            detail=f"Could not extract text from the provided URL. Please paste the text directly. Error: {str(e)}"
        )


async def _resolve_jd(
    job_description: Optional[str] = None,
    jd_file: Optional[UploadFile] = None,
    jd_url: Optional[str] = None
) -> str:
    """Intelligently resolves JD from text, file, or URL."""
    # 1. Check if it's a URL in the text field
    jd_text = (job_description or "").strip()
    if jd_text.startswith(("http://", "https://")) and len(jd_text.split()) == 1:
        return _extract_from_url(jd_text)
    
    # 2. Use URL field if provided
    if jd_url and jd_url.strip():
        return _extract_from_url(jd_url.strip())
    
    # 3. Use File if provided
    if jd_file:
        try:
            content = await jd_file.read()
            return extract_text(content, jd_file.filename)
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"Failed to process job description file: {e}")

    # 4. Use Text if provided
    if jd_text:
        return jd_text
    
    raise HTTPException(
        status_code=422, 
        detail="Job description is missing. Provide 'job_description' (text/url), 'jd_file', or 'jd_url'."
    )


def filename_to_name(filename: str) -> str:
    return Path(filename).stem.replace("_", " ").replace("-", " ").title()


def _collect_files(form) -> List[UploadFile]:
    found, seen = [], set()
    for field in ("resumes", "resume_file", "resume", "file", "files"):
        for item in form.getlist(field):
            if hasattr(item, "read") and id(item) not in seen:
                seen.add(id(item))
                found.append(item)
    for key in form.keys():
        for item in form.getlist(key):
            if hasattr(item, "read") and id(item) not in seen:
                seen.add(id(item))
                found.append(item)
    return found


# ══════════════════════════════════════════════════════════════════
#  ENDPOINTS
# ══════════════════════════════════════════════════════════════════

@app.post(
    "/analyze-resume/",
    tags=["Resume Analysis"],
    summary="Upload one resume → full ATS report",
)
async def analyze_resume_endpoint(
    request:         Request,
    job_description: Optional[str] = Form(None),
    resume_file:     UploadFile    = File(...),
    jd_file:         Optional[UploadFile] = File(None),
    jd_url:          Optional[str] = Form(None),
    candidate_name:  Optional[str] = Form(None),
    user: models.User = Depends(get_current_user),
    db:   Session = Depends(get_db),
):
    try:
        logger.info(f"analyze-resume called by user_id={user.id} file={resume_file.filename}")

        # --- Usage Limit Check ---
        used = _count_usage_this_month(user.id, db)
        if used >= user.monthly_limit:
            raise HTTPException(
                status_code=429,
                detail={
                    "error":              "Monthly request limit reached.",
                    "plan":               user.plan.upper(),
                    "limit":              user.monthly_limit,
                    "used":               used,
                    "remaining":          0,
                    "resets":             "1st of next month (UTC)",
                },
            )

        # Resolve Job Description from any format
        jd_text = await _resolve_jd(job_description, jd_file, jd_url)

        raw_bytes = await resume_file.read()
        if not raw_bytes:
            raise HTTPException(status_code=422, detail="Uploaded file is empty.")

        text   = extract_text(raw_bytes, resume_file.filename)
        name   = candidate_name or filename_to_name(resume_file.filename)

        # Use run_in_threadpool for CPU-bound NLP task
        result = await run_in_threadpool(
            analyze_resume,
            resume_text=text,
            job_description=jd_text,
            candidate_name=name,
        )

        log_usage(user.id, endpoint="/analyze-resume/", db=db)

        used += 1 # Reflect the current request in the response
        result["usage"] = {
            "requests_used":      used,
            "requests_limit":     user.monthly_limit,
            "requests_remaining": max(0, user.monthly_limit - used),
            "plan":               user.plan.upper(),
        }

        logger.info(f"analyze-resume success for user_id={user.id} score={_ats_score_value(result)}")
        return result

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"analyze-resume crashed: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")


@app.post(
    "/bulk-analyze/",
    tags=["Resume Analysis"],
    summary="Upload multiple resumes → ranked ATS list",
    openapi_extra={
        "requestBody": {
            "required": True,
            "content": {
                "multipart/form-data": {
                    "schema": {
                        "type": "object",
                        "required": ["job_description", "resumes"],
                        "properties": {
                            "job_description": {"type": "string"},
                            "resumes": {
                                "type":  "array",
                                "items": {"type": "string", "format": "binary"},
                            },
                        },
                    }
                }
            },
        }
    },
)
async def bulk_analyze(
    request: Request,
    user: models.User = Depends(get_current_user),
    db:   Session      = Depends(get_db),
):
    try:
        logger.info(f"bulk-analyze called by user_id={user.id}")

        try:
            form = await request.form()
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Could not parse form data: {e}")

        # Resolve JD from form fields
        jd_text = await _resolve_jd(
            job_description=form.get("job_description"),
            jd_file=form.get("jd_file"),
            jd_url=form.get("jd_url")
        )

        resume_files = _collect_files(form)
        logger.info(f"bulk-analyze: {len(resume_files)} files received, fields={list(form.keys())}")

        if not resume_files:
            raise HTTPException(
                status_code=422,
                detail={
                    "error":           "No resume files found in the request.",
                    "tip":             "Send files as multipart fields named 'resumes'.",
                    "received_fields": list(form.keys()),
                },
            )

        # Pre-flight usage check
        used      = _count_usage_this_month(user.id, db)
        remaining = user.monthly_limit - used
        if len(resume_files) > remaining:
            raise HTTPException(
                status_code=429,
                detail={
                    "error":               "Not enough requests remaining.",
                    "files_sent":          len(resume_files),
                    "requests_remaining":  remaining,
                    "plan":                user.plan.upper(),
                },
            )

        results: List[dict] = []
        errors:  List[dict] = []

        for resume_file in resume_files:
            fname = getattr(resume_file, "filename", None) or "unknown"
            try:
                raw_bytes = await resume_file.read()
                if not raw_bytes:
                    errors.append({"file": fname, "error": "File is empty."})
                    continue

                text   = extract_text(raw_bytes, fname)
                # Use run_in_threadpool for CPU-bound NLP task
                report = await run_in_threadpool(
                    analyze_resume,
                    resume_text=text,
                    job_description=jd_text,
                    candidate_name=filename_to_name(fname),
                )
                results.append(report)
                
                # Batch usage logging: add but don't commit in loop
                db.add(models.UsageLog(
                    user_id=user.id,
                    request_time=datetime.utcnow(),
                    endpoint="/bulk-analyze/",
                    status="success",
                ))
                
                logger.info(f"bulk-analyze: processed {fname} score={_ats_score_value(report)}")

            except HTTPException as e:
                logger.warning(f"bulk-analyze: skipping {fname} — {e.detail}")
                errors.append({"file": fname, "error": e.detail})
            except Exception as e:
                logger.error(f"bulk-analyze: error on {fname}: {traceback.format_exc()}")
                errors.append({"file": fname, "error": str(e)})

        # Commit all usage logs at once
        if results:
            try:
                db.commit()
            except Exception as e:
                logger.error(f"Failed to commit bulk usage: {e}")
                db.rollback()

        if not results and errors:
            raise HTTPException(
                status_code=422,
                detail={"message": "All files failed to process.", "errors": errors},
            )

        ranked      = sorted(results, key=_ats_score_value, reverse=True)
        shortlisted = [r for r in ranked if _ats_score_value(r) >= 70]
        final_used  = _count_usage_this_month(user.id, db)

        return {
            "summary": {
                "total_uploaded":        len(resume_files),
                "successfully_analyzed": len(ranked),
                "failed":                len(errors),
                "shortlisted_above_70":  len(shortlisted),
                "top_candidate":         ranked[0]["candidate"]["name"] if ranked else None,
                "top_ats_score":         _ats_score_value(ranked[0]) if ranked else None,
            },
            "usage": {
                "requests_used":      final_used,
                "requests_limit":     user.monthly_limit,
                "requests_remaining": max(0, user.monthly_limit - final_used),
                "plan":               user.plan.upper(),
            },
            "ranked_candidates": [
                {
                    "rank":                i + 1,
                    "candidate":           r["candidate"],
                    "scores":              r["scores"],
                    "skill_analysis":      r["skill_analysis"],
                    "jd_keyword_analysis": r["jd_keyword_analysis"],
                    "section_analysis":    r["section_analysis"],
                    "experience":          r["experience"],
                    "education":           r["education"],
                    "job_title_alignment": r["job_title_alignment"],
                    "writing_quality":     r["writing_quality"],
                    "format_and_contact":  r["format_and_contact"],
                    "ai_insights":         r.get("ai_insights"),
                }
                for i, r in enumerate(ranked)
            ],
            "failed_files": errors,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"bulk-analyze crashed: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Bulk analysis failed: {str(e)}")